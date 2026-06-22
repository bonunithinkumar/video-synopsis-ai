# M6 — PDF and DOCX rendering service
# Converts a validated SynopsisInput into a branded PDF (via WeasyPrint)
# or DOCX (via python-docx).

import os
import re
import uuid
from datetime import datetime

from fastapi import HTTPException
from jinja2 import Environment, BaseLoader, select_autoescape

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from weasyprint import HTML

from app.schemas.export import SynopsisInput


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
BRAND_COLOR = "#1a73e8"
BRAND_COLOR_RGB = RGBColor(0x1A, 0x73, 0xE8)
TEXT_COLOR_RGB = RGBColor(0x20, 0x21, 0x24)
MUTED_RGB = RGBColor(0x80, 0x86, 0x8B)

OUTPUT_DIR = os.path.join(os.getcwd(), "generated_exports")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def cleanup_file(path: str) -> None:
    """Background-task callback that deletes a temp export file."""
    try:
        if os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
    return slug[:60] if slug else "synopsis"


# ---------------------------------------------------------------------------
# PDF — Jinja2 template + WeasyPrint
# ---------------------------------------------------------------------------
PDF_TEMPLATE = """
<html>
<head>
<style>
    @page {
        margin: 2.2cm 2cm 2.5cm 2cm;
        @bottom-center {
            content: "Page " counter(page) " of " counter(pages);
            font-size: 9px;
            color: #80868b;
        }
    }
    body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #333333; line-height: 1.6; font-size: 13px; }
    .brand-header { display: flex; justify-content: space-between; align-items: baseline; border-bottom: 3px solid """ + BRAND_COLOR + """; padding-bottom: 6px; margin-bottom: 18px; }
    .brand-header .label { font-size: 11px; letter-spacing: 1px; text-transform: uppercase; color: """ + BRAND_COLOR + """; font-weight: bold; }
    .brand-header .date { font-size: 11px; color: #80868b; }
    h1 { color: #202124; font-size: 22px; margin: 4px 0 14px 0; }
    h2 { color: #202124; margin-top: 24px; font-size: 15px; border-left: 4px solid """ + BRAND_COLOR + """; padding-left: 8px; }
    h3 { color: """ + BRAND_COLOR + """; font-size: 13px; margin: 14px 0 4px 0; }
    .metadata { background-color: #f8f9fa; padding: 12px 14px; border-radius: 6px; margin-bottom: 18px; font-size: 12.5px; border: 1px solid #dadce0; }
    .metadata a { color: """ + BRAND_COLOR + """; text-decoration: none; font-weight: 600; word-break: break-all; }
    .thumb { max-width: 240px; border-radius: 6px; margin-bottom: 14px; display: block; }
    ul { padding-left: 20px; margin-top: 8px; }
    li { margin-bottom: 6px; }
    .topic-card { background: #fafbfc; border: 1px solid #e8eaed; border-radius: 6px; padding: 10px 14px; margin-bottom: 10px; }
    .closing { background: #f0f6ff; border-left: 4px solid """ + BRAND_COLOR + """; padding: 12px 16px; margin-top: 22px; font-style: italic; border-radius: 0 6px 6px 0; }
    .empty { color: #9aa0a6; font-style: italic; font-size: 12px; }
</style>
</head>
<body>
    <div class="brand-header">
        <span class="label">Video Synopsis</span>
        <span class="date">Generated {{ generated_on }}</span>
    </div>

    {% if meta.thumbnail_url %}
    <img class="thumb" src="{{ meta.thumbnail_url }}" />
    {% endif %}

    <h1>{{ meta.title }}</h1>

    <div class="metadata">
        {% if meta.channel_name %}<strong>Channel:</strong> {{ meta.channel_name }}<br/>{% endif %}
        <strong>Source Video:</strong>
        <a href="{{ meta.video_url }}">{{ meta.video_url }}</a>
    </div>

    <h2>Overall Synopsis</h2>
    <p>{{ summary.basic_summary.overall_synopsis }}</p>

    <h2>{{ summary.topics_covered.title }}</h2>
    {% if summary.topics_covered.topics %}
    <ul>{% for t in summary.topics_covered.topics %}<li>{{ t }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No topics provided.</p>{% endif %}

    <h2>Key Insights</h2>
    {% if summary.detailed_summary.key_insights %}
    <ul>{% for i in summary.detailed_summary.key_insights %}<li>{{ i }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No insights provided.</p>{% endif %}

    {% if summary.detailed_summary.topic_breakdown %}
    <h2>Topic Breakdown</h2>
    {% for item in summary.detailed_summary.topic_breakdown %}
    <div class="topic-card">
        <h3>{{ item.topic }}</h3>
        <p>{{ item.explanation }}</p>
    </div>
    {% endfor %}
    {% endif %}

    <h2>Action Items</h2>
    {% if summary.detailed_summary.action_items %}
    <ul>{% for a in summary.detailed_summary.action_items %}<li>{{ a }}</li>{% endfor %}</ul>
    {% else %}<p class="empty">No action items provided.</p>{% endif %}

    <div class="closing">{{ summary.closing_note }}</div>
</body>
</html>
"""

_jinja_env = Environment(loader=BaseLoader(), autoescape=select_autoescape(["html", "xml"]))
_pdf_template = _jinja_env.from_string(PDF_TEMPLATE)


def render_pdf(data: SynopsisInput) -> str:
    """Render a branded PDF from a SynopsisInput. Returns the file path."""
    html_content = _pdf_template.render(
        meta=data.video_metadata,
        summary=data.summary,
        generated_on=datetime.utcnow().strftime("%d %b %Y, %H:%M UTC"),
    )
    file_name = f"{_slugify(data.video_metadata.title)}-{uuid.uuid4().hex[:8]}.pdf"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    try:
        HTML(string=html_content).write_pdf(file_path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")
    return file_path


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def _set_heading_style(paragraph, size=14, color=BRAND_COLOR_RGB):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"


def _add_bullets(doc, items, empty_text="No items provided."):
    if not items:
        p = doc.add_paragraph(empty_text)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x9A, 0xA0, 0xA6)
        return
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def _shade_cell(cell, hex_color):
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


# ---------------------------------------------------------------------------
# DOCX — python-docx
# ---------------------------------------------------------------------------
def render_docx(data: SynopsisInput) -> str:
    """Render a branded DOCX from a SynopsisInput. Returns the file path."""
    try:
        meta = data.video_metadata
        summary = data.summary

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = TEXT_COLOR_RGB

        # Brand strip
        header_p = doc.add_paragraph()
        hr = header_p.add_run("VIDEO SYNOPSIS")
        hr.font.size = Pt(10)
        hr.font.bold = True
        hr.font.color.rgb = BRAND_COLOR_RGB
        dr = header_p.add_run(
            f"   |   Generated {datetime.utcnow().strftime('%d %b %Y, %H:%M UTC')}"
        )
        dr.font.size = Pt(9)
        dr.font.color.rgb = MUTED_RGB

        # Title
        title = doc.add_heading(meta.title, level=1)
        _set_heading_style(title, size=20, color=RGBColor(0x20, 0x21, 0x24))

        # Metadata table
        rows = [("Channel", meta.channel_name)] if meta.channel_name else []
        rows.append(("Source URL", meta.video_url))
        meta_table = doc.add_table(rows=len(rows), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (label, value) in enumerate(rows):
            label_cell, value_cell = meta_table.rows[i].cells
            label_cell.width = Inches(1.3)
            value_cell.width = Inches(5.0)
            _shade_cell(label_cell, "F8F9FA")
            lr = label_cell.paragraphs[0].add_run(label)
            lr.font.bold = True
            lr.font.size = Pt(10)
            vr = value_cell.paragraphs[0].add_run(value)
            vr.font.size = Pt(10)
        doc.add_paragraph()

        # Overall synopsis
        h = doc.add_heading("Overall Synopsis", level=2)
        _set_heading_style(h)
        p = doc.add_paragraph(summary.basic_summary.overall_synopsis)
        p.runs[0].font.size = Pt(11)

        # Topics covered
        h = doc.add_heading(summary.topics_covered.title, level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.topics_covered.topics, "No topics provided.")

        # Key insights
        h = doc.add_heading("Key Insights", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.key_insights, "No insights provided.")

        # Topic breakdown
        if summary.detailed_summary.topic_breakdown:
            h = doc.add_heading("Topic Breakdown", level=2)
            _set_heading_style(h)
            for item in summary.detailed_summary.topic_breakdown:
                sub = doc.add_heading(item.topic, level=3)
                _set_heading_style(sub, size=11, color=BRAND_COLOR_RGB)
                ep = doc.add_paragraph(item.explanation)
                ep.runs[0].font.size = Pt(10.5)

        # Action items
        h = doc.add_heading("Action Items", level=2)
        _set_heading_style(h)
        _add_bullets(doc, summary.detailed_summary.action_items, "No action items provided.")

        # Closing note
        h = doc.add_heading("Closing Note", level=2)
        _set_heading_style(h)
        cp = doc.add_paragraph(summary.closing_note)
        cp.runs[0].font.italic = True

        file_name = f"{_slugify(meta.title)}-{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)
        doc.save(file_path)
        return file_path

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")
